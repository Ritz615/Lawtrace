"""Reports router — PDF and DOCX generation using ReportLab and python-docx."""
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import io

from app.core.database import get_db
from app.core.security import decode_token, oauth2_scheme
from app.models.models import Document, AIReport, User

router = APIRouter()


async def get_current_user(token: str = Depends(oauth2_scheme), db: AsyncSession = Depends(get_db)) -> User:
    payload = decode_token(token)
    result = await db.execute(select(User).where(User.id == payload["sub"]))
    user = result.scalar_one_or_none()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user


@router.get("/{document_id}/pdf")
async def export_pdf(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Generate and stream a PDF report for a document's AI analysis."""
    doc_result = await db.execute(select(Document).where(Document.id == document_id, Document.owner_id == current_user.id))
    doc = doc_result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    report_result = await db.execute(
        select(AIReport).where(AIReport.document_id == document_id).order_by(AIReport.created_at.desc())
    )
    report = report_result.scalar_one_or_none()
    if not report:
        raise HTTPException(404, "No AI report found. Run analysis first.")

    pdf_bytes = _generate_pdf(doc, report)

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{doc.title}_LexAI_Report.pdf"'},
    )


@router.get("/{document_id}/docx")
async def export_docx(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Generate and stream a DOCX report for a document's AI analysis."""
    doc_result = await db.execute(select(Document).where(Document.id == document_id, Document.owner_id == current_user.id))
    doc = doc_result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    report_result = await db.execute(
        select(AIReport).where(AIReport.document_id == document_id).order_by(AIReport.created_at.desc())
    )
    report = report_result.scalar_one_or_none()
    if not report:
        raise HTTPException(404, "No AI report found. Run analysis first.")

    docx_bytes = _generate_docx(doc, report)

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{doc.title}_LexAI_Report.docx"'},
    )


def _generate_pdf(doc, report) -> bytes:
    """Generate a styled PDF using ReportLab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor, black, white
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.units import inch

    buffer = io.BytesIO()
    doc_pdf = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=0.75*inch, rightMargin=0.75*inch)

    PRIMARY = HexColor("#6366f1")
    DARK    = HexColor("#1a1d27")
    MUTED   = HexColor("#94a3b8")
    DANGER  = HexColor("#ef4444")
    SUCCESS = HexColor("#10b981")

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", fontSize=22, textColor=white, spaceAfter=8, fontName="Helvetica-Bold")
    h2 = ParagraphStyle("h2", fontSize=14, textColor=PRIMARY, spaceBefore=16, spaceAfter=6, fontName="Helvetica-Bold")
    body = ParagraphStyle("body", fontSize=10, textColor=DARK, spaceAfter=4, leading=14)
    muted = ParagraphStyle("muted", fontSize=9, textColor=MUTED, spaceAfter=4)

    risk_color = {
        "low": SUCCESS, "medium": HexColor("#f59e0b"), "high": DANGER, "critical": HexColor("#dc2626")
    }.get(str(report.risk_level.value if report.risk_level else "").lower(), MUTED)

    story = []

    # Header block
    header_data = [[
        Paragraph("LexAI – Contract Analysis Report", ParagraphStyle("hdr", fontSize=20, textColor=white, fontName="Helvetica-Bold")),
        Paragraph(f"Risk Score: {report.risk_score or 'N/A'}/100", ParagraphStyle("rs", fontSize=28, textColor=risk_color, fontName="Helvetica-Bold", alignment=2)),
    ]]
    header_table = Table(header_data, colWidths=[4.5*inch, 2.5*inch])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), DARK),
        ("PADDING", (0, 0), (-1, -1), 16),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [DARK]),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 12))

    # Details table
    story.append(Paragraph("Contract Details", h2))
    details = [
        ["Document", doc.original_filename or doc.title],
        ["Contract Type", str(report.contract_type.value) if report.contract_type else "Unknown"],
        ["Parties", ", ".join(report.parties) if report.parties else "N/A"],
        ["Effective Date", str(report.effective_date or "N/A")],
        ["Expiration Date", str(report.expiration_date or "N/A")],
        ["Risk Level", str(report.risk_level.value).upper() if report.risk_level else "N/A"],
        ["AI Confidence", f"{int((report.ai_confidence or 0) * 100)}%"],
    ]
    dt = Table([[Paragraph(k, muted), Paragraph(v, body)] for k, v in details], colWidths=[2*inch, 5*inch])
    dt.setStyle(TableStyle([
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [HexColor("#f8fafc"), white]),
        ("PADDING", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
    ]))
    story.append(dt)

    # Summary
    if report.executive_summary:
        story.append(Paragraph("Executive Summary", h2))
        story.append(Paragraph(report.executive_summary, body))

    # Key Risks
    if report.key_risks:
        story.append(Paragraph("Key Risks", h2))
        for risk in report.key_risks:
            story.append(Paragraph(f"• {risk}", ParagraphStyle("risk", fontSize=10, textColor=DANGER, leading=14, spaceAfter=4)))

    # Recommendations
    if report.recommendations:
        story.append(Paragraph("Recommendations", h2))
        for i, rec in enumerate(report.recommendations, 1):
            story.append(Paragraph(f"{i}. {rec}", ParagraphStyle("rec", fontSize=10, textColor=SUCCESS, leading=14, spaceAfter=4)))

    # Missing clauses
    if report.missing_clauses:
        story.append(Paragraph("Missing Clauses", h2))
        story.append(Paragraph(", ".join(report.missing_clauses), ParagraphStyle("mc", fontSize=10, textColor=HexColor("#f59e0b"), leading=14)))

    story.append(Spacer(1, 24))
    story.append(HRFlowable(width="100%", color=HexColor("#e2e8f0")))
    story.append(Paragraph("Generated by LexAI – Enterprise AI Contract Management Platform", muted))

    doc_pdf.build(story)
    buffer.seek(0)
    return buffer.read()


def _generate_docx(doc, report) -> bytes:
    """Generate a styled DOCX report using python-docx."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = DocxDocument()

    # Title
    title = d.add_heading("LexAI – Contract Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d.add_paragraph(f"Document: {doc.original_filename or doc.title}")
    d.add_paragraph(f"Risk Score: {report.risk_score or 'N/A'}/100  |  Risk Level: {str(report.risk_level.value).upper() if report.risk_level else 'N/A'}")
    d.add_paragraph(f"Contract Type: {str(report.contract_type.value) if report.contract_type else 'Unknown'}")
    d.add_paragraph(f"Parties: {', '.join(report.parties) if report.parties else 'N/A'}")
    d.add_paragraph(f"Effective: {report.effective_date or 'N/A'}  |  Expires: {report.expiration_date or 'N/A'}")
    d.add_paragraph("")

    if report.executive_summary:
        d.add_heading("Executive Summary", 1)
        d.add_paragraph(report.executive_summary)

    if report.key_risks:
        d.add_heading("Key Risks", 1)
        for risk in report.key_risks:
            p = d.add_paragraph(f"• {risk}", style="List Bullet")
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)

    if report.recommendations:
        d.add_heading("Recommendations", 1)
        for i, rec in enumerate(report.recommendations, 1):
            d.add_paragraph(f"{i}. {rec}", style="List Number")

    if report.missing_clauses:
        d.add_heading("Missing Clauses", 1)
        for clause in report.missing_clauses:
            d.add_paragraph(f"• {clause}", style="List Bullet")

    d.add_paragraph("\nGenerated by LexAI – Enterprise AI Contract Management Platform")

    buffer = io.BytesIO()
    d.save(buffer)
    buffer.seek(0)
    return buffer.read()
